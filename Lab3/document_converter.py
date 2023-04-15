import os
import string
import random

from selenium import webdriver
from selenium.webdriver.chrome.service import Service
from selenium.webdriver.chrome.options import Options
from selenium.webdriver.common.by import By
from selenium.webdriver.remote.webelement import WebElement
from webdriver_manager.chrome import ChromeDriverManager


# constants
notebook_path = '/home/choleraplague/university/DataScience/TextAnalysisWithPython/TextAnalysis/Lab3/'
notebook_name = 'lab3'
out_name = 'АТП_Лаб3_ІП11_Панченко.docx'
front_path = 'front.docx'
back_path = 'back.docx'

os.system(f'jupyter nbconvert --to html {notebook_name}.ipynb')

# init driver
options = Options()
options.add_argument("--headless")
options.add_argument("--window-size=850,1440")
options.add_experimental_option('detach', True)
driver = webdriver.Chrome(
    service=Service(ChromeDriverManager().install()), options=options)

def get_random_string(length: int) -> str:
    return ''.join(random.choice(string.ascii_lowercase) for i in range(length))

def parse_notebook(driver, path) -> list:
    index = 1
    tree = []
    driver.get(url=path)
    html_cells = driver.find_elements(By.CLASS_NAME, 'jp-Cell')
    for c in html_cells:
        class_name = c.get_attribute('class')
        class_name_parts = class_name.split(' ')
        if 'jp-MarkdownCell' in class_name_parts:
            rendered = c.find_element(By.CLASS_NAME, 'jp-RenderedMarkdown')
            heading = rendered.find_element(By.XPATH, '*')
            tag_attr = heading.tag_name
            if tag_attr == 'h1':
                tree.append(
                    {'text': heading.text, 'children': []})
            elif tag_attr == 'h2':
                tree[-1]['children'].append(
                    {'text': heading.text, 'children': []})
            elif tag_attr == 'h3':
                tree[-1]['children'][-1]['children'].append({'text': heading.text, 'picture': None, 'name': None})
            elif tag_attr == 'p':
                tree[-1]['children'][-1]['children'][-1]['name'] = heading.text
        elif 'jp-CodeCell' in class_name_parts:
            name = str(index) + '.png'
            index += 1
            c.screenshot(name)
            tree[-1]['children'][-1]['children'][-1]['picture'] = name
    return tree

tree = parse_notebook(driver, 'file://' + notebook_path + notebook_name + '.html')

from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.shared import RGBColor
from PIL import Image
document = Document()

# def add_heading_one_style(document) -> str:
#     styles = document.styles
#     name = 'New Heading 1'
#     new_heading_style = styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)
#     new_heading_style.base_style = styles['Heading 1']
#     font = new_heading_style.font
#     font.name = 'Times New Roman'
#     font.size = Pt(14)
#     new_heading_style.paragraph_format.first_line_indent = Inches(0)
#     new_heading_style.font.color.rgb = RGBColor(0, 0, 0)
#     new_heading_style.font.all_caps = True
#     new_heading_style.font.bold = False
#     return name

# def add_heading_two_style(document) -> str:
#     styles = document.styles
#     name = 'New Heading 2'
#     new_heading_style = styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)
#     new_heading_style.base_style = styles['Heading 2']
#     font = new_heading_style.font
#     font.name = 'Times New Roman'
#     font.size = Pt(14)
#     new_heading_style.paragraph_format.first_line_indent = Inches(0)
#     new_heading_style.font.color.rgb = RGBColor(0, 0, 0)
#     new_heading_style.font.bold = False
#     return name

new_heading_one_style_name = 'Heading 1'# add_heading_one_style(document)
new_heading_two_style_name = 'Heading 2' #add_heading_two_style(document)

def update_section_styles(document):
    for section in document.sections:
        section.left_margin = Inches(0.98)
        section.right_margin = Inches(0.59)
        section.top_margin = Inches(0.49)
        section.bottom_margin = Inches(0.59)

# def update_paragraph_format_styles(document):
#     paragraph_format = document.styles['Default Paragraph Style'].paragraph_format
#     paragraph_format.line_spacing = 1.5
#     paragraph_format.first_line_indent = Inches(0.492)

def update_picture_para_style(picture_para, name):
    picture_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    picture_run = picture_para.add_run()
    im = Image.open(name)
    # in pixels
    w, h = im.size
    # to inches 
    w /= 136
    h /= 136
    picture = picture_run.add_picture(name, width=Inches(w), height=Inches(h))

def update_run_font(run):
    font = run.font
    font.size = Pt(14)
    font.name = 'Times New Roman'
    font.color.rgb = RGBColor(0, 0, 0)

def update_para_ident(para):
    para.paragraph_format.line_spacing = 1.5
    para.paragraph_format.first_line_indent = Inches(0.492)
    para.paragraph_format.right_indent = Inches(0.309)

def update_document(document):
    for head_one in tree:
        head_one_parts = head_one['text'].split(' ')
        head_one_number = head_one_parts[0]
        heading_one = document.add_paragraph(' '.join(head_one_parts[1:]), style=new_heading_one_style_name)
        heading_one.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        for head_two in head_one['children']:
            heading_two = document.add_paragraph(style=new_heading_two_style_name)
            run = heading_two.add_run(' '.join(head_two['text'].split(' ')[1:]) + '\n')
            heading_two.is_linked_to_previous = True
            update_run_font(run)
            for element in head_two['children']:
                para_text = document.add_paragraph()
                update_para_ident(para_text)
                ex_run = para_text.add_run(element['text'])
                update_run_font(ex_run)
                if element['picture']:
                    document.add_paragraph(' ')
                    picture_number = element['picture'].split('.')[0]
                    picture_text = element['name']
                    picture_para = document.add_paragraph()
                    update_para_ident(picture_para)
                    update_picture_para_style(picture_para, element['picture'])
                    run_text = f'\n\nРисунок {head_one_number}.{picture_number} - {picture_text}\n'
                    run = picture_para.add_run(run_text)
                    update_run_font(run)

update_section_styles(document)
#update_paragraph_format_styles(document)
update_document(document)

document.save('test.docx')
os.system('rm *.png')

from docxcompose.composer import Composer


def combine_all_docx(combined_name, filename_master, files_list):
    number_of_sections=len(files_list)
    master = Document(filename_master)
    composer = Composer(master)
    for i in range(0, number_of_sections):
        doc_temp = Document(files_list[i])
        composer.append(doc_temp)
    composer.save(combined_name)

combine_all_docx(out_name, front_path, ['test.docx', back_path])


    
    


        


        
